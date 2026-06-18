#!/usr/bin/env python3
"""Extract academic paper text into Markdown with best-effort page markers."""

from __future__ import annotations

import argparse
import json
from pathlib import Path


def read_text(path: Path) -> tuple[str, list[str]]:
    return path.read_text(encoding="utf-8", errors="replace"), []


def read_pdf(path: Path) -> tuple[str, list[str]]:
    notes: list[str] = []
    pages: list[str] = []

    try:
        import pdfplumber  # type: ignore

        with pdfplumber.open(path) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text(layout=True) or ""
                pages.append(f"[Page {index}]\n\n{text.strip()}")
        return "\n\n".join(pages).strip() + "\n", notes
    except Exception as exc:
        notes.append(f"pdfplumber extraction unavailable or failed: {exc}")

    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            pages.append(f"[Page {index}]\n\n{text.strip()}")
        return "\n\n".join(pages).strip() + "\n", notes
    except Exception as exc:
        notes.append(f"pypdf extraction unavailable or failed: {exc}")

    raise RuntimeError("Could not extract PDF text. Install pdfplumber or pypdf, or provide a text/Markdown source.")


def read_docx(path: Path) -> tuple[str, list[str]]:
    notes = ["DOCX extraction cannot reliably recover original printed page numbers."]
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise RuntimeError("Could not import python-docx. Install python-docx or provide a text/Markdown source.") from exc

    document = docx.Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n[Table {table_index}]")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n\n".join(parts).strip() + "\n", notes


def extract(path: Path) -> tuple[str, dict]:
    suffix = path.suffix.lower()
    readers = {
        ".txt": read_text,
        ".md": read_text,
        ".markdown": read_text,
        ".tex": read_text,
        ".latex": read_text,
        ".pdf": read_pdf,
        ".docx": read_docx,
    }
    if suffix not in readers:
        raise RuntimeError(f"Unsupported source format: {suffix}")

    text, notes = readers[suffix](path)
    metadata = {
        "source_path": str(path),
        "source_format": suffix.lstrip("."),
        "extraction_notes": notes,
        "characters": len(text),
    }
    return text, metadata


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract paper text to source.md.")
    parser.add_argument("source", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()

    args.out.mkdir(parents=True, exist_ok=True)
    text, metadata = extract(args.source)

    (args.out / "source.md").write_text(text, encoding="utf-8")
    (args.out / "source_metadata.json").write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    print(f"Wrote {args.out / 'source.md'}")
    if metadata["extraction_notes"]:
        print("Notes:")
        for note in metadata["extraction_notes"]:
            print(f"- {note}")


if __name__ == "__main__":
    main()

